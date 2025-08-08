from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML
from io import BytesIO
from docx import Document
from datetime import datetime
import os

TEMPLATES = os.path.join(os.path.dirname(__file__), "templates")

def render_html(context: dict, company_name: str, brand: dict | None = None) -> str:
    env = Environment(loader=FileSystemLoader(TEMPLATES), autoescape=select_autoescape())
    tpl = env.get_template("report.html")
    brand = brand or {"primary":"#16BAC5","accent":"#1A1D2E","logo_b64":None}
    return tpl.render(company_name=company_name, brand=brand, generated_at=datetime.utcnow(), **context)

def render_pdf_bytes(context: dict, company_name: str, brand: dict | None = None) -> bytes:
    html = render_html(context, company_name, brand=brand)
    pdf_io = BytesIO()
    HTML(string=html, base_url=TEMPLATES).write_pdf(pdf_io)
    return pdf_io.getvalue()

def render_docx_bytes(context: dict, company_name: str) -> bytes:
    doc = Document()
    doc.add_heading(f"{company_name} — Agentic Maturity Assessment", 0)
    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_heading("Summary", level=1)
    overall = context["scores"]["overall"]
    doc.add_paragraph(f"Overall Weighted Score: {overall['weighted_score']} (Level {overall['level']})")
    doc.add_heading("Domain Scores", level=1)
    for d, s in context["scores"]["domains"].items():
        doc.add_paragraph(f"- {d}: {s['weighted_score']} (Level {s['level']})")
    doc.add_heading("Roadmap (Top Actions)", level=1)
    for item in context["roadmap"][:15]:
        doc.add_paragraph(f"[{item['risk']}] {item['domain']} - {item['question']} ⇒ {item['recommended_action']}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
